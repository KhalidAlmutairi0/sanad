"""Upload sanitizer extraction — runs INSIDE the bubblewrap sandbox (no network).

Input: a raw file (pdf/docx/txt) bind-mounted read-only at a fixed path.
Output: plain UTF-8/NFC text on stdout. Nothing else survives — no macros, scripts,
embedded objects, or formatting. Prompt-injection strings in the text are NOT removed here
(impossible to do reliably); they are neutralized at the prompt layer by untrusted-data
tagging. This process has no network and cannot import anything network-bound.

Exit codes: 0 ok; 2 unsupported type; 3 extraction error; 4 empty result.
"""
from __future__ import annotations

import os
import pathlib
import sys
import unicodedata

SUPPORTED = {".pdf", ".docx", ".txt"}

# spec #3: OCR pages whose mean Tesseract word-confidence falls below this are flagged so the
# reviewer knows the text came from a poor scan. Env-overridable (mirrors OCR_MIN_CONFIDENCE).
_OCR_MIN_CONFIDENCE = float(os.environ.get("OCR_MIN_CONFIDENCE", "60"))


def _mean_confidence(confidences: list) -> float:
    """Mean of valid per-word confidences (Tesseract uses -1 for non-text boxes). Empty -> 0."""
    vals = [float(c) for c in confidences if c not in (None, "", "-1") and float(c) >= 0]
    return sum(vals) / len(vals) if vals else 0.0


def _clean(text: str) -> str:
    # NFC normalize; strip control chars except tab/newline; collapse trailing whitespace.
    text = unicodedata.normalize("NFC", text)
    out = []
    for ch in text:
        if ch in ("\t", "\n"):
            out.append(ch)
            continue
        if unicodedata.category(ch).startswith("C"):
            continue  # drop control / format chars (defensive)
        out.append(ch)
    lines = [ln.rstrip() for ln in "".join(out).splitlines()]
    return "\n".join(lines).strip()


# A digital PDF yields a real text layer; a scanned (image) PDF yields almost nothing from
# pypdf. Below this many non-space characters we treat the PDF as scanned and OCR it.
_SCANNED_TEXT_THRESHOLD = 80
_OCR_DPI = 200
_OCR_MAX_PAGES = 40  # bound work under the sandbox timeout/rlimits


def _ocr_pdf(path: pathlib.Path) -> str:
    """OCR a scanned PDF with Tesseract (ara+eng), fully offline. Runs inside the no-network
    sandbox — tesseract + poppler need no network, so the containment guarantee holds.

    Also captures per-page mean word-confidence (image_to_data) and flags the document when any
    page is below the threshold (spec #3), so downstream can mark it 'verify against original'."""
    from pdf2image import convert_from_path
    import pytesseract

    images = convert_from_path(str(path), dpi=_OCR_DPI, first_page=1, last_page=_OCR_MAX_PAGES)
    pages = []
    low_confidence = False
    for img in images:
        pages.append(pytesseract.image_to_string(img, lang="ara+eng"))
        data = pytesseract.image_to_data(img, lang="ara+eng", output_type=pytesseract.Output.DICT)
        if _mean_confidence(data.get("conf", [])) < _OCR_MIN_CONFIDENCE:
            low_confidence = True
    # Signal to the wrapper that OCR was used (kept off stdout, which is the clean text).
    sys.stderr.write("OCR_USED\n")
    if low_confidence:
        sys.stderr.write("LOW_OCR_CONFIDENCE\n")
    return "\n".join(pages)


def _extract_pdf(path: pathlib.Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    if len(text.replace(" ", "").strip()) < _SCANNED_TEXT_THRESHOLD:
        # No usable text layer → scanned document. Fall back to OCR.
        return _ocr_pdf(path)
    return text


def _extract_docx(path: pathlib.Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_txt(path: pathlib.Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        sys.stderr.write("usage: extract.py <input_path>\n")
        return 3
    path = pathlib.Path(argv[1])
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        sys.stderr.write(f"unsupported_file_type: {suffix}\n")
        return 2
    try:
        if suffix == ".pdf":
            raw = _extract_pdf(path)
        elif suffix == ".docx":
            raw = _extract_docx(path)
        else:
            raw = _extract_txt(path)
    except Exception as exc:  # noqa: BLE001 — any parser failure is an extraction error
        sys.stderr.write(f"extraction_error: {type(exc).__name__}\n")
        return 3

    text = _clean(raw)
    if not text:
        sys.stderr.write("empty_result\n")
        return 4
    sys.stdout.write(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
