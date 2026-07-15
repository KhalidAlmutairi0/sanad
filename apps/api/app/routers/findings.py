from __future__ import annotations

import datetime as dt
import uuid

from fastapi import APIRouter, Depends, Query
from sqlalchemy import case, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.deps import get_current_user, get_session, require_roles
from app.core.errors import SanadError
from app.core.storage import get_minio, presigned_get
from app.models import Clause, Contract, Finding, Regulation, RegulationVersion, User
from app.schemas.findings import (
    Citation,
    ExplainResponse,
    FindingItem,
    FindingList,
    KitExportRequest,
    KitExportResponse,
    KitResponse,
    RadarKiller,
    RadarResponse,
    ReviewRequest,
    ReviewResponse,
)
from app.services.analysis.kit import generate_explanation, generate_kit
from app.services.audit import write_audit
from app.services.scoring import compute_readiness_score
from app.services.scoring.score import compute_radar

_settings = get_settings()

router = APIRouter(tags=["findings"])

_SEVERITY_RANK = case(
    {"critical": 0, "high": 1, "medium": 2, "low": 3},
    value=Finding.severity,
    else_=4,
)


def _citation(rv: RegulationVersion, code: str) -> Citation:
    return Citation(
        regulation_version_id=rv.id,
        regulation_code=code,
        article_ref=rv.article_ref,
        article_text_ar=rv.article_text_ar,
        source_url=rv.source_url,
        effective_date=rv.effective_date,
        verification_tier=rv.verification_tier,
    )


def _item(f: Finding, rv: RegulationVersion, code: str) -> FindingItem:
    return FindingItem(
        id=f.id,
        clause_id=f.clause_id,
        title_ar=f.title_ar,
        title_en=f.title_en,
        explanation_ar=f.explanation_ar,
        explanation_en=f.explanation_en,
        severity=f.severity,
        category=f.category,
        violation_cost_ar=f.violation_cost_ar,
        violation_cost_min=float(f.violation_cost_min) if f.violation_cost_min is not None else None,
        violation_cost_max=float(f.violation_cost_max) if f.violation_cost_max is not None else None,
        confidence_tier=f.confidence_tier,
        review_status=f.review_status,
        citation=_citation(rv, code),
    )


@router.get("/contracts/{contract_id}/findings", response_model=FindingList)
async def list_findings(
    contract_id: uuid.UUID,
    status: str | None = Query(default=None, pattern="^(pending|accepted|rejected)$"),
    severity: str | None = Query(default=None, pattern="^(critical|high|medium|low)$"),
    _: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> FindingList:
    stmt = (
        select(Finding, RegulationVersion, Regulation.code)
        .join(RegulationVersion, RegulationVersion.id == Finding.regulation_version_id)
        .join(Regulation, Regulation.id == RegulationVersion.regulation_id)
        .where(Finding.contract_id == contract_id)
        .order_by(_SEVERITY_RANK, Finding.created_at)
    )
    if status:
        stmt = stmt.where(Finding.review_status == status)
    if severity:
        stmt = stmt.where(Finding.severity == severity)

    rows = (await session.execute(stmt)).all()
    return FindingList(items=[_item(f, rv, code) for f, rv, code in rows])


@router.get("/findings/{finding_id}/explain", response_model=ExplainResponse)
async def explain_finding(
    finding_id: uuid.UUID,
    _: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> ExplainResponse:
    row = (
        await session.execute(
            select(Finding, RegulationVersion, Regulation.code)
            .join(RegulationVersion, RegulationVersion.id == Finding.regulation_version_id)
            .join(Regulation, Regulation.id == RegulationVersion.regulation_id)
            .where(Finding.id == finding_id)
        )
    ).first()
    if not row:
        raise SanadError("not_found")
    f, rv, code = row
    # Regenerate a plain-language explanation strictly from the cited article.
    gen = await generate_explanation(
        article_text=rv.article_text_ar, code=code, article_ref=rv.article_ref, title=f.title_ar
    )
    return ExplainResponse(
        explanation_ar=gen.get("explanation_ar") or f.explanation_ar,
        explanation_en=gen.get("explanation_en") or f.explanation_en,
        citation=_citation(rv, code),
    )


@router.post("/findings/{finding_id}/review", response_model=ReviewResponse)
async def review_finding(
    finding_id: uuid.UUID,
    body: ReviewRequest,
    user: User = Depends(require_roles("reviewer", "admin")),
    session: AsyncSession = Depends(get_session),
) -> ReviewResponse:
    if body.decision not in ("accepted", "rejected"):
        raise SanadError("validation_failed")
    finding = await session.get(Finding, finding_id)
    if not finding:
        raise SanadError("not_found")
    # A finding is decided once; re-deciding is a conflict, not a silent overwrite.
    if finding.review_status != "pending":
        raise SanadError("review_conflict")

    finding.review_status = body.decision
    finding.reviewed_by = user.id
    finding.reviewed_at = dt.datetime.now(dt.timezone.utc)
    await write_audit(
        session, actor=str(user.id), action="finding_reviewed",
        target=str(finding.contract_id), verdict="n-a",
        detail={"finding_id": str(finding_id), "decision": body.decision},
    )

    # Reviewed-only score recompute (writes its own score_computed audit).
    await compute_readiness_score(session, finding.contract_id)

    # Contract is 'reviewed' once nothing is left pending.
    pending = (
        await session.execute(
            select(func.count()).select_from(Finding).where(
                Finding.contract_id == finding.contract_id,
                Finding.review_status == "pending",
            )
        )
    ).scalar_one()
    contract = await session.get(Contract, finding.contract_id)
    if contract is not None:
        contract.status = "reviewed" if pending == 0 else "reviewing"

    await session.commit()
    return ReviewResponse(
        id=finding.id, review_status=finding.review_status, reviewed_at=finding.reviewed_at
    )


@router.get("/contracts/{contract_id}/radar", response_model=RadarResponse)
async def contract_radar(
    contract_id: uuid.UUID,
    _: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> RadarResponse:
    radar = await compute_radar(session, contract_id)
    killers = []
    for f in radar["killers"]:
        row = (
            await session.execute(
                select(RegulationVersion, Regulation.code)
                .join(Regulation, Regulation.id == RegulationVersion.regulation_id)
                .where(RegulationVersion.id == f.regulation_version_id)
            )
        ).first()
        if not row:
            continue
        rv, code = row
        killers.append(
            RadarKiller(finding_id=f.id, title_ar=f.title_ar, severity=f.severity, citation=_citation(rv, code))
        )
    return RadarResponse(verdict=radar["verdict"], killers=killers)


async def _finding_row(session: AsyncSession, finding_id: uuid.UUID):
    return (
        await session.execute(
            select(Finding, RegulationVersion, Regulation.code)
            .join(RegulationVersion, RegulationVersion.id == Finding.regulation_version_id)
            .join(Regulation, Regulation.id == RegulationVersion.regulation_id)
            .where(Finding.id == finding_id)
        )
    ).first()


async def _clause_text(session: AsyncSession, clause_id: uuid.UUID | None) -> str:
    if not clause_id:
        return ""
    clause = await session.get(Clause, clause_id)
    return (clause.text_ar or clause.text_en or "") if clause else ""


@router.get("/findings/{finding_id}/kit", response_model=KitResponse)
async def finding_kit(
    finding_id: uuid.UUID,
    _: User = Depends(require_roles("reviewer", "admin")),
    session: AsyncSession = Depends(get_session),
) -> KitResponse:
    row = await _finding_row(session, finding_id)
    if not row:
        raise SanadError("not_found")
    f, rv, code = row
    if not (f.review_status == "accepted" and f.severity in ("critical", "high")):
        raise SanadError("validation_failed",
                         "ملف التفاوض يُنشأ فقط للملاحظات المقبولة عالية الخطورة",
                         "Negotiation Kit is only for accepted high or critical findings")
    kit = await generate_kit(
        clause_text=await _clause_text(session, f.clause_id),
        article_text=rv.article_text_ar, code=code, article_ref=rv.article_ref,
    )
    return KitResponse(
        redrafted_clause_ar=kit.get("redrafted_clause_ar", ""),
        redrafted_clause_en=kit.get("redrafted_clause_en", ""),
        justification_letter_ar=kit.get("justification_letter_ar", ""),
        justification_letter_en=kit.get("justification_letter_en", ""),
        citation=_citation(rv, code),
    )


def _build_annex(kit: dict, code: str, article_ref: str, fmt: str) -> tuple[bytes, str, str]:
    """Return (bytes, content_type, extension). Stacked bilingual annex."""
    if fmt == "docx":
        import io

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading("ملحق التفاوض · Negotiation Annex", level=1)
        doc.add_paragraph(f"{code} {article_ref}")
        for ar_head, ar_key, en_head, en_key in [
            ("الصياغة المقترحة", "redrafted_clause_ar", "Redrafted clause", "redrafted_clause_en"),
            ("خطاب التبرير", "justification_letter_ar", "Justification", "justification_letter_en"),
        ]:
            h = doc.add_heading(ar_head, level=2); h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = doc.add_paragraph(kit.get(ar_key, "")); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_heading(en_head, level=2)
            doc.add_paragraph(kit.get(en_key, ""))
        buf = io.BytesIO(); doc.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

    # pdf not rendered on-box; return a plain-text annex (stacked bilingual).
    text = (
        f"ملحق التفاوض · {code} {article_ref}\n\n"
        f"الصياغة المقترحة:\n{kit.get('redrafted_clause_ar','')}\n\n"
        f"Redrafted clause:\n{kit.get('redrafted_clause_en','')}\n\n"
        f"خطاب التبرير:\n{kit.get('justification_letter_ar','')}\n\n"
        f"Justification:\n{kit.get('justification_letter_en','')}\n"
    )
    return text.encode("utf-8"), "text/plain; charset=utf-8", "txt"


@router.post("/findings/{finding_id}/kit/export", response_model=KitExportResponse)
async def finding_kit_export(
    finding_id: uuid.UUID,
    body: KitExportRequest,
    user: User = Depends(require_roles("reviewer", "admin")),
    session: AsyncSession = Depends(get_session),
) -> KitExportResponse:
    import io

    row = await _finding_row(session, finding_id)
    if not row:
        raise SanadError("not_found")
    f, rv, code = row
    kit = await generate_kit(
        clause_text=await _clause_text(session, f.clause_id),
        article_text=rv.article_text_ar, code=code, article_ref=rv.article_ref,
    )
    fmt = body.format if body.format in ("docx", "pdf") else "docx"
    data, content_type, ext = _build_annex(kit, code, rv.article_ref, fmt)
    key = f"{f.contract_id}/kits/{finding_id}.{ext}"
    get_minio().put_object(
        _settings.bucket_sanitized, key, io.BytesIO(data), length=len(data), content_type=content_type
    )
    await write_audit(
        session, actor=str(user.id), action="kit_exported", target=str(finding_id),
        verdict="n-a", detail={"format": fmt, "key": key},
    )
    await session.commit()
    return KitExportResponse(download_url=presigned_get(_settings.bucket_sanitized, key))
